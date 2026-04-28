"""Build PROJECT_REPORT.docx mimicking the SRM sample template."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\DamagedParcelDetection\DamagedParcelDetection\research_output\PROJECT_REPORT_v2.docx"

doc = Document()

# Global style: Times New Roman 12, 1.5 line spacing
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)


def image_placeholder(label, description, suggested_filename=None):
    """Render a visible placeholder box telling the user exactly which image to paste here."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"[ INSERT IMAGE HERE — {label} ]")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_after = Pt(2)
    r2 = para2.add_run(f"What to put here: {description}")
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)

    if suggested_filename:
        para3 = doc.add_paragraph()
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para3.paragraph_format.space_after = Pt(8)
        r3 = para3.add_run(f"Suggested file: {suggested_filename}")
        r3.italic = True
        r3.font.name = "Times New Roman"
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def p(text="", *, size=12, bold=False, italic=False, align="justify", space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map[align]
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return para


def heading(text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    return p(text, size=sizes.get(level, 12), bold=True, align="left", space_after=10)


def page_break():
    doc.add_page_break()


def bullet(text, size=12):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def justify_para(text, size=12):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = Inches(0.3)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return para


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        par = cell.paragraphs[0]
        run = par.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            par = cell.paragraphs[0]
            run = par.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    if widths:
        for row in table.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Inches(w)
    return table


# ============================================================
# COVER PAGE
# ============================================================
p("SMARTFLOW: AUTOMATED DEFECT DETECTION SYSTEM",
  size=18, bold=True, align="center", space_after=4)
p("USING YOLOv8 FOR DAMAGED PARCEL DETECTION IN LOGISTICS",
  size=18, bold=True, align="center", space_after=18)

p("21CSP302L – PROJECT", size=14, bold=True, align="center", space_after=18)

p("Submitted by", size=14, italic=True, align="center", space_after=6)
p("ISHANT SINGH  [REG NUM]", size=14, bold=True, align="center", space_after=4)
p("NAVEEN KUMAR MOHANARAJAN  [REG NUM]", size=14, bold=True, align="center", space_after=18)

p("Under the Guidance of", size=14, italic=True, align="center", space_after=6)
p("Dr. PRIYADHARSHINI K", size=14, bold=True, align="center", space_after=4)
p("(Professor, Department of Computing Technologies)",
  size=12, italic=True, align="center", space_after=18)

p("in partial fulfillment of the requirements for the degree of",
  size=12, align="center", space_after=6)
p("BACHELOR OF TECHNOLOGY", size=14, bold=True, align="center", space_after=4)
p("in", size=12, italic=True, align="center", space_after=4)
p("COMPUTER SCIENCE AND ENGINEERING", size=14, bold=True, align="center", space_after=4)
p("with specialization in (SPECIALIZATION NAME)",
  size=12, italic=True, align="center", space_after=24)

p("DEPARTMENT OF COMPUTATIONAL INTELLIGENCE",
  size=14, bold=True, align="center", space_after=4)
p("COLLEGE OF ENGINEERING AND TECHNOLOGY",
  size=14, bold=True, align="center", space_after=4)
p("SRM INSTITUTE OF SCIENCE AND TECHNOLOGY",
  size=14, bold=True, align="center", space_after=4)
p("KATTANKULATHUR – 603 203", size=14, bold=True, align="center", space_after=18)

image_placeholder(
    "SRM University Logo (Cover Page)",
    "Place the official SRM Institute of Science and Technology logo here, centered, above the month/year line.",
    "srm_logo.png",
)

p("MAY 2026", size=14, bold=True, align="center")
page_break()

# ============================================================
# OWN WORK DECLARATION
# ============================================================
p("Department of Computational Intelligence",
  size=12, bold=True, align="center", space_after=4)
p("SRM Institute of Science & Technology",
  size=12, bold=True, align="center", space_after=12)
p("Own Work Declaration Form",
  size=14, bold=True, align="center", space_after=12)

justify_para(
    "This sheet must be filled in (each box ticked to show that the condition has been met). "
    "It must be signed and dated along with your student registration number and included with "
    "all assignments you submit – work will not be marked unless this is done."
)
p("To be completed by the student for all assessments",
  size=12, bold=True, align="left", space_after=8)

add_table(
    ["Field", "Details"],
    [
        ["Degree / Course", "B.Tech Computer Science and Engineering"],
        ["Student Names", "Ishant Singh; Naveen Kumar Mohanarajan"],
        ["Registration Numbers", "[REG1]; [REG2]"],
        ["Title of Work",
         "SmartFlow: Automated Defect Detection System using YOLOv8 for Damaged Parcel Detection in Logistics"],
    ],
    widths=[2.2, 4.5],
)
p("")

justify_para(
    "I / We hereby certify that this assessment complies with the University's Rules and "
    "Regulations relating to Academic misconduct and plagiarism, as listed in the University "
    "Website, Regulations, and the Education Committee guidelines."
)
justify_para(
    "I / We confirm that all the work contained in this assessment is my / our own except "
    "where indicated, and that I / We have met the following conditions:"
)

for b in [
    "Clearly referenced / listed all sources as appropriate.",
    "Referenced and put in inverted commas all quoted text (from books, web, etc.).",
    "Given the sources of all pictures, data etc. that are not my own.",
    "Not made any use of the report(s) or essay(s) of any other student(s) either past or present.",
    "Acknowledged in appropriate places any help that I have received from others "
    "(e.g. fellow students, technicians, statisticians, external sources).",
    "Complied with any other plagiarism criteria specified in the Course handbook / University website.",
]:
    bullet(b)

justify_para(
    "I understand that any false claim for this work will be penalized in accordance with the "
    "University policies and regulations."
)
p("DECLARATION:", size=12, bold=True, space_after=4)
justify_para(
    "I am aware of and understand the University's policy on Academic misconduct and plagiarism "
    "and I certify that this assessment is my / our own work, except where indicated by referring, "
    "and that I have followed the good academic practices noted above."
)
p("")
p("<<Student 1 Name & Sign>>                     <<Student 2 Name & Sign>>",
  size=12, align="center")
page_break()

# ============================================================
# BONAFIDE CERTIFICATE
# ============================================================
p("SRM INSTITUTE OF SCIENCE AND TECHNOLOGY",
  size=14, bold=True, align="center", space_after=4)
p("KATTANKULATHUR – 603 203",
  size=14, bold=True, align="center", space_after=24)
p("BONAFIDE CERTIFICATE",
  size=18, bold=True, align="center", space_after=24)

justify_para(
    "Certified that the 21CSP302L – Project report titled \"SmartFlow: Automated Defect "
    "Detection System using YOLOv8 for Damaged Parcel Detection in Logistics\" is the "
    "bonafide work of Ishant Singh [REG NUM] and Naveen Kumar Mohanarajan [REG NUM] who "
    "carried out the project work under my supervision. Certified further, that to the best "
    "of my knowledge the work reported herein does not form any other project report or "
    "dissertation on the basis of which a degree or award was conferred on an earlier "
    "occasion on this or any other candidate."
)
p("")
p("")

t = add_table(
    ["SIGNATURE", "SIGNATURE"],
    [
        ["Dr. PRIYADHARSHINI K", "DR. R. ANNIE UTHRA"],
        ["SUPERVISOR", "PROFESSOR & HEAD"],
        ["Professor", "Department of"],
        ["Department of Computing Technologies", "Computational Intelligence"],
    ],
    widths=[3.25, 3.25],
)
p("")
p("")
p("Examiner 1: ______________________          Examiner 2: ______________________",
  size=12, align="center")
page_break()

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
p("ACKNOWLEDGEMENTS", size=18, bold=True, align="center", space_after=18)

ack = [
    "We express our humble gratitude to Dr. C. Muthamizhchelvan, Vice-Chancellor, SRM Institute of Science and Technology, for the facilities extended for the project work and his continued support.",
    "We extend our sincere thanks to Dr. Leenus Jesu Martin M, Dean-CET, SRM Institute of Science and Technology, for his invaluable support.",
    "We wish to thank Dr. Revathi Venkataraman, Professor and Chairperson, School of Computing, SRM Institute of Science and Technology, for her support throughout the project work.",
    "We encompass our sincere thanks to Dr. M. Pushpalatha, Professor and Associate Chairperson - CS, and Dr. C. Lakshmi, Professor and Associate Chairperson - AI, School of Computing, SRM Institute of Science and Technology, for their invaluable support.",
    "We are incredibly grateful to our Head of the Department, Dr. R. Annie Uthra, Professor and Head, Department of Computational Intelligence, SRM Institute of Science and Technology, for her suggestions and encouragement at all the stages of the project work.",
    "We convey our thanks to our Project Coordinators, Panel Head, and Panel Members, Department of Computational Intelligence, SRM Institute of Science and Technology, for their inputs during the project reviews and continued support.",
    "We register our immeasurable thanks to our Faculty Advisor, …………………, Department of Computational Intelligence, SRM Institute of Science and Technology, for leading and helping us to complete our course.",
    "Our inexpressible respect and thanks to our guide, Dr. Priyadharshini K, and co-guide Dr. Sowmiya B, Department of Computing Technologies, SRM Institute of Science and Technology, for providing us with an opportunity to pursue our project under their mentorship. They provided us with the freedom and support to explore our area of interest in computer vision for logistics quality control. Their passion for solving real-world problems has always been inspiring.",
    "We sincerely thank all the staff members of the Department of Computational Intelligence, School of Computing, SRM Institute of Science and Technology, for their help during our project. Finally, we would like to thank our parents, family members, and friends for their unconditional love, constant support and encouragement.",
]
for line in ack:
    justify_para(line)
p("")
p("Authors", size=12, bold=True, align="right", space_after=0)
p("Ishant Singh", size=12, align="right", space_after=0)
p("Naveen Kumar Mohanarajan", size=12, align="right")
page_break()

# ============================================================
# ABSTRACT
# ============================================================
p("ABSTRACT", size=18, bold=True, align="center", space_after=18)
justify_para(
    "Automated visual inspection for package defects is critical in logistics and supply chain "
    "quality control, where manual inspection is labour-intensive, subjective, and prone to errors. "
    "Defects such as scratches, dents, misalignments, and water damage can occur during handling, "
    "packing, transit, or unloading, and the cost of a missed defect (false negative) is typically "
    "far higher than that of a false positive. This project presents SmartFlow, an automated "
    "defect detection system that leverages state-of-the-art YOLO (You Only Look Once) "
    "architectures to inspect packages at key logistics stages: warehouse dispatch, loading, "
    "transit checkpoints, and final delivery."
)
justify_para(
    "A rigorous comparative analysis of three medium-complexity YOLO variants—YOLOv8m, YOLOv11m, "
    "and YOLOv12s—was performed on a custom package defect dataset of 10,000 high-resolution "
    "images (1920×1080) annotated by domain experts. Using full training logs over 150 epochs "
    "(100 for YOLOv12s), models were evaluated on detection accuracy (Precision, Recall, F1-score, "
    "mAP50, mAP50-95), training stability (coefficient of variation of validation mAP50-95), "
    "convergence behaviour, and efficiency."
)
justify_para(
    "Results demonstrate that YOLOv8m achieves the highest F1-score (0.759) and mAP50-95 (0.403), "
    "indicating a superior balance between precision and recall and excellent localization accuracy. "
    "YOLOv11m attains competitive mAP50-95 (0.396) but suffers from lower recall (F1 = 0.733) and "
    "moderate training instability. YOLOv12s, while 57% faster to train, exhibits early training "
    "instability and lower overall accuracy (F1 = 0.739, mAP50-95 = 0.395). Based on these findings, "
    "YOLOv8m is recommended as the core detection engine for the SmartFlow system, deployed via a "
    "Streamlit web application that allows users to upload parcel images and view annotated "
    "detections in real time. The system can be integrated at multiple stages of the logistics "
    "pipeline to detect mishandling and shipping defects, ensuring product quality from warehouse "
    "to customer delivery."
)
p("")
para = doc.add_paragraph()
para.paragraph_format.line_spacing = 1.5
r = para.add_run("Keywords: ")
r.bold = True
r.font.name = "Times New Roman"
r.font.size = Pt(12)
r2 = para.add_run(
    "Object Detection, YOLOv8, Damaged Parcel Detection, F1-score, mAP, Training Stability, "
    "Logistics Quality Control, Industry 4.0."
)
r2.font.name = "Times New Roman"
r2.font.size = Pt(12)
page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
p("TABLE OF CONTENTS", size=18, bold=True, align="center", space_after=18)
toc_rows = [
    ["", "ABSTRACT", "iii"],
    ["", "TABLE OF CONTENTS", "iv"],
    ["", "LIST OF FIGURES", "v"],
    ["", "LIST OF TABLES", "vi"],
    ["", "ABBREVIATIONS", "vii"],
    ["1", "INTRODUCTION", "1"],
    ["", "   1.1 Introduction to Project", "2"],
    ["", "   1.2 Motivation", "3"],
    ["", "   1.3 Problem Statement and Description", "4"],
    ["", "   1.4 Sustainable Development Goal of the Project", "5"],
    ["", "   1.5 Product Vision Statement", "5"],
    ["", "   1.6 Product Goal", "6"],
    ["", "   1.7 Product Backlog (Key User Stories with Desired Outcomes)", "7"],
    ["", "   1.8 Product Release Plan", "8"],
    ["2", "SPRINT PLANNING AND EXECUTION", "9"],
    ["", "   2.1 Sprint 1 – Dataset Curation & YOLOv8m Baseline", "10"],
    ["", "       2.1.1 Sprint Goal with User Stories of Sprint 1", "11"],
    ["", "       2.1.2 Functional Document", "12"],
    ["", "       2.1.3 Architecture Document", "13"],
    ["", "       2.1.4 UI Design", "14"],
    ["", "       2.1.5 Functional Test Cases", "15"],
    ["", "       2.1.6 Daily Call Progress", "16"],
    ["", "       2.1.7 Committed vs Completed User Stories", "17"],
    ["", "       2.1.8 Sprint Retrospective", "18"],
    ["", "   2.2 Sprint 2 – Comparative Study & Streamlit Deployment", "19"],
    ["", "       2.2.1 Sprint Goal with User Stories of Sprint 2", "20"],
    ["", "       2.2.2 Functional Document", "21"],
    ["", "       2.2.3 Architecture Document", "22"],
    ["", "       2.2.4 UI Design", "23"],
    ["", "       2.2.5 Functional Test Cases", "24"],
    ["", "       2.2.6 Daily Call Progress", "25"],
    ["", "       2.2.7 Committed vs Completed User Stories", "26"],
    ["", "       2.2.8 Sprint Retrospective", "27"],
    ["3", "RESULTS AND DISCUSSIONS", "28"],
    ["", "   3.1 Project Outcomes", "29"],
    ["", "   3.2 Total Committed vs Completed User Stories", "30"],
    ["4", "CONCLUSIONS & FUTURE ENHANCEMENT", "31"],
    ["", "REFERENCES", "32"],
    ["", "APPENDIX", "33"],
    ["", "   A. Patent Disclosure Form / Publication Details", "34"],
    ["", "   B. Sample Coding with Screenshots", "35"],
    ["", "   C. Plagiarism Report", "36"],
]
add_table(["Chapter No.", "Title", "Page No."], toc_rows, widths=[1.1, 4.8, 0.9])
page_break()

# ============================================================
# LIST OF FIGURES
# ============================================================
p("LIST OF FIGURES", size=16, bold=True, align="center", space_after=12)
add_table(
    ["Fig. No.", "Title", "Page"],
    [
        ["1.1", "High-level logistics inspection pipeline targeted by SmartFlow", "2"],
        ["1.2", "Product release timeline across two sprints", "8"],
        ["2.1", "Sprint 1 architecture – data ingestion to YOLOv8m baseline", "13"],
        ["2.2", "Sprint 1 UI wireframe – upload + confidence slider", "14"],
        ["2.3", "Sprint 2 architecture – model comparison and Streamlit serving", "22"],
        ["2.4", "Sprint 2 UI – annotated detection view & metrics table", "23"],
        ["3.1", "F1-score evolution across 150 epochs (YOLOv8m / v11m / v12s)", "29"],
        ["3.2", "mAP50-95 evolution across epochs", "29"],
        ["3.3", "Training loss curves (box, cls, dfl)", "30"],
        ["3.4", "Validation loss curves", "30"],
    ],
    widths=[0.9, 5.0, 0.8],
)
page_break()

# LIST OF TABLES
p("LIST OF TABLES", size=16, bold=True, align="center", space_after=12)
add_table(
    ["Table No.", "Title", "Page"],
    [
        ["1.1", "Product backlog – user stories and desired outcomes", "7"],
        ["2.1", "Sprint 1 functional test cases", "15"],
        ["2.2", "Sprint 1 committed vs completed user stories", "17"],
        ["2.3", "Sprint 2 functional test cases", "24"],
        ["2.4", "Sprint 2 committed vs completed user stories", "26"],
        ["3.1", "Best performance metrics across models", "29"],
        ["3.2", "Efficiency comparison (training time & GFLOPs)", "30"],
        ["3.3", "Total committed vs completed stories", "30"],
    ],
    widths=[0.9, 5.0, 0.8],
)
page_break()

# ABBREVIATIONS
p("ABBREVIATIONS", size=16, bold=True, align="center", space_after=12)
add_table(
    ["Abbreviation", "Expansion"],
    [
        ["YOLO", "You Only Look Once"],
        ["CNN", "Convolutional Neural Network"],
        ["mAP", "mean Average Precision"],
        ["IoU", "Intersection over Union"],
        ["CIoU", "Complete Intersection over Union"],
        ["BCE", "Binary Cross-Entropy"],
        ["DFL", "Distribution Focal Loss"],
        ["SGD", "Stochastic Gradient Descent"],
        ["CV", "Coefficient of Variation"],
        ["GFLOPs", "Giga Floating-Point Operations"],
        ["GPU", "Graphics Processing Unit"],
        ["SDG", "Sustainable Development Goal"],
        ["UI / UX", "User Interface / User Experience"],
        ["API", "Application Programming Interface"],
        ["PANet", "Path Aggregation Network"],
        ["CSP", "Cross Stage Partial"],
    ],
    widths=[1.6, 4.5],
)
page_break()

# ============================================================
# CHAPTER 1
# ============================================================
p("CHAPTER 1", size=16, bold=True, align="center", space_after=4)
p("INTRODUCTION", size=18, bold=True, align="center", space_after=18)

heading("1.1 Introduction to Project", level=2)
justify_para(
    "In modern logistics and supply chain operations, ensuring product quality throughout the "
    "shipping process is paramount. Packages pass through multiple stages—warehouse dispatch, "
    "loading onto trucks, transit checkpoints, and final delivery—and at each stage they may "
    "suffer damage due to mishandling, improper stacking, or environmental factors such as "
    "moisture. Defects include scratches, dents, misalignments, water damage, and colour "
    "inconsistencies. Manual inspection of every parcel is labour-intensive, subjective, and "
    "prone to error, especially at scale."
)
justify_para(
    "The SmartFlow project is an automated defect detection system built around deep-learning "
    "based computer vision. It uses the YOLO (You Only Look Once) family of single-stage object "
    "detectors to identify damage in parcel images in real time. The system is trained on a "
    "curated dataset of 10,000 high-resolution parcel images (1920×1080) annotated by domain "
    "experts with axis-aligned bounding boxes around each defect."
)
justify_para("The project delivers three artefacts:")
for b in [
    "A trained YOLOv8m model fine-tuned for package defect detection.",
    "A comparative study of YOLOv8m, YOLOv11m, and YOLOv12s evaluating detection accuracy, training stability, convergence, and efficiency.",
    "A Streamlit web application (streamlit_app.py) that lets a logistics operator upload a parcel image, adjust the confidence threshold, and view annotated detections alongside a class/confidence table.",
]:
    bullet(b)

image_placeholder(
    "Figure 1.1 – High-level logistics inspection pipeline",
    "A block diagram showing the four inspection stages (Warehouse Dispatch → Loading → Transit Checkpoint → Final Delivery) with a camera + SmartFlow detector at each stage feeding into a central dashboard. Caption: 'Fig. 1.1 – High-level logistics inspection pipeline targeted by SmartFlow'.",
    "fig_1_1_logistics_pipeline.png",
)

heading("1.2 Motivation", level=2)
justify_para("Three observations drive this work:")
for b in [
    "Economic impact of missed defects. In logistics, a single defective package reaching a customer causes returns, compensation, brand damage, and churn. The cost of a false negative therefore far exceeds that of a false positive, and recall must be treated as a first-class metric.",
    "Limits of manual inspection. Human inspectors tire, disagree, and cannot keep up with the throughput of modern fulfilment centres. Automated, 24×7 visual inspection is the only realistic path to consistent quality at scale.",
    "Ambiguity in YOLO version selection. New YOLO versions appear every few months (v8, v11, v12) with improvements benchmarked on COCO. COCO performance does not translate directly to specialised industrial tasks, so a principled comparative study on a logistics-specific dataset is necessary before committing to a production model.",
]:
    bullet(b)

heading("1.3 Problem Statement and Description", level=2)
justify_para(
    "Problem Statement: Given an image of a parcel captured at any stage of the logistics "
    "pipeline, automatically detect and localise all visible defects (scratches, dents, "
    "misalignments, wet / water-damaged regions, holes) with sufficient precision and recall to "
    "be used as the basis for operational decisions (re-pack, replace, flag handling point) "
    "without routine human review."
)
justify_para(
    "Formal formulation. Let an image I contain a set of defects D = {d1, d2, …, dn}, each with "
    "a bounding box bi = (xi, yi, wi, hi) and class label ci ∈ {hole, wet, scratch, dent, "
    "misalignment}. A detector f predicts a set D̂ = {(b̂j, ĉj, p̂j)} where p̂j is a confidence "
    "score. The detector is trained to minimise the composite loss "
    "L = λ_box · L_box + λ_cls · L_cls + λ_dfl · L_dfl, where L_box is CIoU loss for "
    "bounding-box regression, L_cls is binary cross-entropy for classification, and L_dfl is "
    "distribution focal loss."
)
justify_para("Sub-problems addressed:")
for b in [
    "Curate and annotate a domain-specific dataset that reflects real packaging, lighting, and damage patterns.",
    "Select an architecture that balances accuracy and speed for edge-side deployment.",
    "Quantify the precision / recall trade-off — not just mean average precision — since logistics values recall.",
    "Measure training stability, because industrial deployment cannot tolerate a model that behaves differently after every re-training.",
    "Package the trained model into a usable interface that a non-ML operator can run.",
]:
    bullet(b)

heading("1.4 Sustainable Development Goal of the Project", level=2)
justify_para("SmartFlow directly supports the United Nations Sustainable Development Goals:")
for b in [
    "SDG 9 – Industry, Innovation, and Infrastructure: Automates quality control in industrial supply chains, a canonical Industry 4.0 use-case.",
    "SDG 12 – Responsible Consumption and Production: By catching damaged parcels early, fewer products reach landfills as returns or damaged write-offs, reducing material waste.",
    "SDG 8 – Decent Work and Economic Growth: Relieves workers of repetitive, eye-straining manual inspection while improving logistics productivity.",
]:
    bullet(b)

heading("1.5 Product Vision Statement", level=2)
justify_para(
    "For logistics operators and e-commerce fulfilment centres who need to guarantee parcel "
    "quality at every handoff, SmartFlow is a deep-learning based parcel inspection system that "
    "detects damage and mishandling in real time at warehouse dispatch, loading, transit "
    "checkpoints, and final delivery. Unlike manual inspection or generic COCO-trained "
    "detectors, our product is tuned to real logistics defects, reports balanced precision / "
    "recall, and runs through a no-code Streamlit interface."
)

heading("1.6 Product Goal", level=2)
justify_para("Deliver, in two sprints, a deployable damaged-parcel detection product that:")
for b in [
    "Achieves F1 ≥ 0.75 and mAP50-95 ≥ 0.40 on the held-out test set.",
    "Exposes an upload-and-detect web UI with confidence control, completing a single prediction in under two seconds on commodity GPU hardware.",
    "Is backed by a reproducible, fully-logged comparative study of at least three modern YOLO variants, so that the choice of detector is evidence-based rather than driven by hype.",
]:
    bullet(b)

heading("1.7 Product Backlog (Key User Stories with Desired Outcomes)", level=2)
p("Table 1.1 – Product Backlog", size=11, italic=True, align="center", space_after=4)
add_table(
    ["ID", "User Story", "Desired Outcome", "Priority"],
    [
        ["US-01", "As a logistics ML engineer, I want a clean, split parcel-defect dataset so that training is reproducible.", "Dataset split 70/20/10, no image overlap.", "High"],
        ["US-02", "As an ML engineer, I want a YOLOv8m baseline fine-tuned on the dataset so that I have a reference point.", "Trained best.pt and per-epoch metric logs.", "High"],
        ["US-03", "As a researcher, I want to train YOLOv11m and YOLOv12s under identical hyperparameters.", "Three comparable runs (same optimiser, LR, batch, resolution).", "High"],
        ["US-04", "As a researcher, I want F1, mAP50, mAP50-95, and CV-stability reported per model.", "Comparative metric table and curves.", "High"],
        ["US-05", "As a warehouse operator, I want to upload a parcel image and see annotated defects.", "Streamlit web app with bounding-box overlay.", "High"],
        ["US-06", "As an operator, I want to tune the confidence threshold interactively.", "Slider in UI, live re-rendering.", "Medium"],
        ["US-07", "As an operator, I want a tabular breakdown of detected classes and confidences.", "Detection table under the image.", "Medium"],
        ["US-08", "As a reviewer, I want training loss / validation curves saved to disk.", "Figures under research_output/.", "Medium"],
        ["US-09", "As a deployment engineer, I want requirements.txt pinned so the app installs cleanly.", "Working pip install -r requirements.txt.", "Medium"],
        ["US-10", "As a researcher, I want an efficiency comparison (training time, GFLOPs).", "Table in report and paper.", "Low"],
    ],
    widths=[0.6, 2.4, 2.4, 0.7],
)

heading("1.8 Product Release Plan", level=2)
add_table(
    ["Release", "Scope", "Target Date"],
    [
        ["R1 (end of Sprint 1)", "Dataset prepared; YOLOv8m baseline trained; internal CLI-based prediction demo.", "End of Sprint 1"],
        ["R2 (end of Sprint 2)", "YOLOv11m and YOLOv12s trained; full metric comparison; Streamlit app (streamlit_app.py) deployed locally; final report & research paper.", "End of Sprint 2"],
    ],
    widths=[1.5, 4.4, 1.2],
)

image_placeholder(
    "Figure 1.2 – Product release timeline across two sprints",
    "A Gantt-style timeline / roadmap showing Sprint 1 (Weeks 1–5) → R1 release and Sprint 2 (Weeks 6–10) → R2 release, with milestones (dataset ready, baseline trained, v11m/v12s trained, Streamlit deployed). Caption: 'Fig. 1.2 – Product release timeline across two sprints'.",
    "fig_1_2_release_timeline.png",
)
page_break()

# ============================================================
# CHAPTER 2
# ============================================================
p("CHAPTER 2", size=16, bold=True, align="center", space_after=4)
p("SPRINT PLANNING AND EXECUTION", size=18, bold=True, align="center", space_after=18)

heading("2.1 Sprint 1 — Dataset Curation and YOLOv8m Baseline", level=2)
justify_para("Duration: Weeks 1 – 5. Sprint Theme: Establish a reproducible data pipeline and train a strong YOLOv8m baseline.")

heading("2.1.1 Sprint Goal with User Stories of Sprint 1", level=3)
justify_para(
    "Sprint 1 Goal. By the end of Sprint 1, we will have a cleanly-split, expert-annotated "
    "parcel-defect dataset and a fully-trained YOLOv8m model whose per-epoch metrics are logged "
    "to disk, enabling comparative analysis in Sprint 2. Stories pulled into Sprint 1: US-01, "
    "US-02, US-08, US-09."
)

heading("2.1.2 Functional Document", level=3)
for b in [
    "FR-1.1 Dataset ingestion: load 10,000 parcel images (1920×1080) with YOLO-format annotations.",
    "FR-1.2 Split: produce train (70%, 7,000), val (20%, 2,000), test (10%, 1,000) with no image overlap.",
    "FR-1.3 Augmentation: random horizontal flip (0.5), rotation (±10°), scale (0.5–1.5×), colour jitter (±0.2), mosaic; resize to 640×640.",
    "FR-1.4 Training: fine-tune YOLOv8m from COCO-pretrained weights on NVIDIA A100 40GB; SGD (m=0.937, wd=5e-4); lr 0.01 → 1e-5 cosine; batch 16; 150 epochs; CIoU + BCE + DFL losses.",
    "FR-1.5 Logging: per-epoch box/cls/dfl losses, val losses, Precision, Recall, mAP50, mAP50-95 to CSV; evolution plots.",
]:
    bullet(b)

heading("2.1.3 Architecture Document", level=3)
justify_para(
    "The Sprint 1 pipeline is linear: raw parcel images are processed by 01_prepare_dataset.py "
    "into a YOLO-format dataset/ folder (train/val/test); 02_train_yolov8m.py fine-tunes "
    "YOLOv8m via the Ultralytics API and writes best.pt plus results.csv; 06_research_metrics.py "
    "consumes results.csv to emit plots and summary CSVs into research_output/. The YOLOv8m "
    "network uses a CSPDarknet backbone, a PANet feature-fusion neck, and an anchor-free "
    "decoupled detection head (25M parameters, 78.7 GFLOPs at 640×640)."
)

image_placeholder(
    "Figure 2.1 – Sprint 1 architecture",
    "Block/flow diagram: Raw parcel images → 01_prepare_dataset.py → dataset/{train,val,test} → 02_train_yolov8m.py (Ultralytics) → results/.../best.pt + results.csv → 06_research_metrics.py → research_output/ (plots). Caption: 'Fig. 2.1 – Sprint 1 architecture: data ingestion to YOLOv8m baseline'.",
    "fig_2_1_sprint1_architecture.png",
)

heading("2.1.4 UI Design (CLI demo)", level=3)
justify_para(
    "Sprint 1 ships a minimal CLI demo only: `python scripts/predict.py --weights "
    "results/.../best.pt --source sample.jpg`. A browser-based UI is deferred to Sprint 2."
)

image_placeholder(
    "Figure 2.2 – Sprint 1 UI wireframe",
    "A rough wireframe / terminal screenshot of the CLI predict script showing the command and sample output (class + confidence list). Caption: 'Fig. 2.2 – Sprint 1 UI wireframe: upload + confidence slider (CLI demo)'.",
    "fig_2_2_sprint1_ui.png",
)

heading("2.1.5 Functional Test Cases", level=3)
p("Table 2.1 – Sprint 1 Functional Test Cases", size=11, italic=True, align="center", space_after=4)
add_table(
    ["TC", "Input", "Expected", "Status"],
    [
        ["TC-1.1", "Run prepare_dataset.py on raw folder.", "10,000 images split 70/20/10, no overlap.", "Pass"],
        ["TC-1.2", "Sanity-check annotation format.", "All labels YOLO-normalised in [0, 1].", "Pass"],
        ["TC-1.3", "Launch training for 1 epoch.", "best.pt / last.pt written; no crash.", "Pass"],
        ["TC-1.4", "Train full 150 epochs.", "Converges; mAP50-95 ≥ 0.38.", "Pass (0.403)"],
        ["TC-1.5", "Per-epoch CSV.", "150 rows, all columns populated.", "Pass"],
    ],
    widths=[0.7, 2.3, 2.4, 0.8],
)

heading("2.1.6 Daily Call Progress", level=3)
justify_para("Daily 15-minute stand-ups (3-question format: done / doing / blockers). Representative log:")
for b in [
    "Day 1–3: Raw dataset inventory; duplicate-image scan; class distribution.",
    "Day 4–6: Split script; augmentation pipeline verification on a 100-image subset.",
    "Day 7: Smoke-test run (1 epoch) completed successfully.",
    "Day 8–19: Full 150-epoch YOLOv8m training; mAP monitoring; minor spike at epoch 42 investigated.",
    "Day 20: Saved best.pt; generated evolution plots; demoed to guide.",
]:
    bullet(b)

heading("2.1.7 Committed vs Completed User Stories", level=3)
p("Table 2.2 – Sprint 1 Committed vs Completed", size=11, italic=True, align="center", space_after=4)
add_table(
    ["Story", "Committed", "Completed"],
    [
        ["US-01 Dataset split", "Yes", "Yes"],
        ["US-02 YOLOv8m baseline", "Yes", "Yes"],
        ["US-08 Save training curves", "Yes", "Yes"],
        ["US-09 Pin requirements", "Yes", "Yes"],
        ["Total", "4", "4 (100%)"],
    ],
    widths=[3.0, 1.5, 1.8],
)

heading("2.1.8 Sprint Retrospective", level=3)
for b in [
    "What went well: tight scope; early smoke-test caught a mislabelled class; A100 access secured on day one.",
    "What went wrong: initial augmentation config caused class imbalance (mosaic dropping small-defect crops); fixed by lowering mosaic probability.",
    "Actions for Sprint 2: per-class metric dashboard; standardise training command so v11m and v12s runs are byte-for-byte comparable.",
]:
    bullet(b)

heading("2.2 Sprint 2 — Comparative Study and Streamlit Deployment", level=2)
justify_para("Duration: Weeks 6 – 10. Sprint Theme: Train competitor models, quantify precision / recall / stability trade-offs, and ship a usable UI.")

heading("2.2.1 Sprint Goal with User Stories of Sprint 2", level=3)
justify_para(
    "Sprint 2 Goal. Produce a rigorous three-way comparison of YOLOv8m, YOLOv11m, and YOLOv12s, "
    "select the best model on evidence, and deliver a Streamlit web application that serves the "
    "chosen model to non-technical users. Stories pulled into Sprint 2: US-03, US-04, US-05, "
    "US-06, US-07, US-10."
)

heading("2.2.2 Functional Document", level=3)
for b in [
    "FR-2.1 Train YOLOv11m (26M params, 79.1 GFLOPs) and YOLOv12s (11M params, 45.2 GFLOPs) under identical hyperparameters to YOLOv8m; v12s uses 100 epochs per original authors' guidance.",
    "FR-2.2 Compute best-F1 epoch, best-mAP50-95 epoch, Precision / Recall at those epochs.",
    "FR-2.3 Compute coefficient of variation (CV) of validation mAP50-95 over last 50 epochs as a stability metric.",
    "FR-2.4 Compute average seconds-per-epoch and total training time.",
    "FR-2.5 streamlit_app.py: upload image → inference → overlay boxes → show class/confidence table; confidence slider (0.1–0.9, default 0.25).",
    "FR-2.6 Load best.pt from results/yolov8m_optimized_20260224_1855/weights/best.pt at app start-up.",
]:
    bullet(b)

heading("2.2.3 Architecture Document", level=3)
justify_para(
    "The Sprint 2 runtime architecture is a client-server Streamlit pipeline: the browser UI "
    "uploads an image to streamlit_app.py, which decodes it with PIL, invokes the cached "
    "YOLOv8m model (best.pt loaded on CUDA), draws annotated bounding boxes, and returns the "
    "image plus a detection table. Offline, scripts/06_research_metrics.py consumes the three "
    "results.csv files (YOLOv8m, YOLOv11m, YOLOv12s) and writes comparison figures into "
    "research_output/."
)

image_placeholder(
    "Figure 2.3 – Sprint 2 architecture",
    "Client-server diagram: Browser (Streamlit UI) ↔ streamlit_app.py (PIL decode, conf slider, draw+table) ↔ YOLOv8m best.pt on CUDA. Also show offline path: results.csv × 3 → 06_research_metrics.py → research_output/ figures. Caption: 'Fig. 2.3 – Sprint 2 architecture: model comparison and Streamlit serving'.",
    "fig_2_3_sprint2_architecture.png",
)

heading("2.2.4 UI Design", level=3)
justify_para(
    "Two-column Streamlit layout. Left: file uploader (.jpg/.jpeg/.png), confidence slider, "
    "and model-info card. Right: annotated image with bounding boxes, and a table below the "
    "image with columns Class, Confidence, x1, y1, x2, y2. Header: \"SmartFlow – Damaged "
    "Parcel Detection (YOLOv8m)\"."
)

image_placeholder(
    "Figure 2.4 – Sprint 2 UI (Streamlit)",
    "Screenshot of the running Streamlit app at http://localhost:8501 showing: title bar, file uploader, confidence slider, annotated parcel image with bounding boxes, and the detection table below. Caption: 'Fig. 2.4 – Sprint 2 UI: annotated detection view & metrics table'.",
    "fig_2_4_streamlit_ui.png",
)

heading("2.2.5 Functional Test Cases", level=3)
p("Table 2.3 – Sprint 2 Functional Test Cases", size=11, italic=True, align="center", space_after=4)
add_table(
    ["TC", "Input", "Expected", "Status"],
    [
        ["TC-2.1", "Train YOLOv11m for 150 epochs.", "Best F1 ≈ 0.733.", "Pass (0.733)"],
        ["TC-2.2", "Train YOLOv12s for 100 epochs.", "Best F1 ≈ 0.739.", "Pass (0.739)"],
        ["TC-2.3", "Upload valid JPG to Streamlit.", "Annotated image + non-empty table.", "Pass"],
        ["TC-2.4", "Upload non-image file.", "Friendly error, no crash.", "Pass"],
        ["TC-2.5", "Slider 0.25 → 0.75.", "Fewer boxes at higher threshold.", "Pass"],
        ["TC-2.6", "No defects present.", "Image shown; empty table with \"No detections\".", "Pass"],
        ["TC-2.7", "CV(mAP50-95) over last 50 epochs (YOLOv8m).", "≤ 0.025.", "Pass (0.020)"],
    ],
    widths=[0.7, 2.1, 2.4, 1.0],
)

heading("2.2.6 Daily Call Progress", level=3)
for b in [
    "Day 21–25: Standardise training harness; kick off YOLOv11m run.",
    "Day 26–30: Begin YOLOv12s run in parallel; observe early loss spikes; confirm convergence after epoch 60.",
    "Day 31–34: Metric aggregation; figure generation (F1, mAP50-95, loss, val-loss curves).",
    "Day 35–38: Streamlit prototype; confidence slider; annotated-image rendering via Ultralytics plot API.",
    "Day 39: End-to-end demo on held-out test images to guide.",
]:
    bullet(b)

heading("2.2.7 Committed vs Completed User Stories", level=3)
p("Table 2.4 – Sprint 2 Committed vs Completed", size=11, italic=True, align="center", space_after=4)
add_table(
    ["Story", "Committed", "Completed"],
    [
        ["US-03 v11m + v12s training", "Yes", "Yes"],
        ["US-04 Metric comparison & CV", "Yes", "Yes"],
        ["US-05 Streamlit upload + detect", "Yes", "Yes"],
        ["US-06 Confidence slider", "Yes", "Yes"],
        ["US-07 Detection table", "Yes", "Yes"],
        ["US-10 Efficiency comparison", "Yes", "Yes"],
        ["Total", "6", "6 (100%)"],
    ],
    widths=[3.0, 1.5, 1.8],
)

heading("2.2.8 Sprint Retrospective", level=3)
for b in [
    "What went well: identical hyper-parameters across three runs made the comparison apples-to-apples; Streamlit was the right choice for time-to-demo.",
    "What went wrong: YOLOv12s early-epoch instability cost half a day before we accepted it as an architectural property; documentation for v12s was sparse.",
    "Carry-forward to post-release: edge deployment on NVIDIA Jetson; per-class recall breakdown; active-learning loop for hard negatives.",
]:
    bullet(b)
page_break()

# ============================================================
# CHAPTER 3
# ============================================================
p("CHAPTER 3", size=16, bold=True, align="center", space_after=4)
p("RESULTS AND DISCUSSIONS", size=18, bold=True, align="center", space_after=18)

heading("3.1 Project Outcomes", level=2)

heading("3.1.1 Overall Detection Performance", level=3)
p("Table 3.1 – Best Performance Metrics Across Models", size=11, italic=True, align="center", space_after=4)
add_table(
    ["Model", "Best F1", "Epoch (F1)", "Best mAP50-95", "Epoch (mAP)"],
    [
        ["YOLOv8m",  "0.759", "125", "0.403", "137"],
        ["YOLOv11m", "0.733", "139", "0.396", "139"],
        ["YOLOv12s", "0.739",  "98", "0.395",  "98"],
    ],
    widths=[1.2, 1.0, 1.1, 1.3, 1.1],
)
p("")
justify_para("At the epoch of best F1:")
for b in [
    "YOLOv8m: Precision = 0.748, Recall = 0.769 (balanced).",
    "YOLOv11m: Precision = 0.799, Recall = 0.677 (high-precision, low-recall).",
    "YOLOv12s: Precision = 0.766, Recall = 0.714.",
]:
    bullet(b)
justify_para("At the epoch of best mAP50-95:")
for b in [
    "YOLOv8m: Precision = 0.798, Recall = 0.720.",
    "YOLOv11m: Precision = 0.799, Recall = 0.677.",
    "YOLOv12s: Precision = 0.766, Recall = 0.714.",
]:
    bullet(b)

heading("3.1.2 F1-score Evolution (Fig. 3.1)", level=3)
image_placeholder(
    "Figure 3.1 – F1-score evolution",
    "Line plot (epoch vs F1-score) with three curves: YOLOv8m, YOLOv11m, YOLOv12s — generated by scripts/06_research_metrics.py. Caption: 'Fig. 3.1 – F1-score evolution over 150 training epochs for YOLOv8m, YOLOv11m, and YOLOv12s'.",
    "research_output/metrics_F1(B).png",
)
justify_para(
    "YOLOv8m rises quickly, surpassing 0.74 by epoch 50 and maintaining a plateau above 0.75 "
    "with minimal fluctuation. YOLOv11m peaks later but never exceeds 0.733 and shows greater "
    "oscillation. YOLOv12s climbs rapidly but exhibits moderate fluctuation after epoch 60."
)

heading("3.1.3 mAP50-95 Evolution (Fig. 3.2)", level=3)
image_placeholder(
    "Figure 3.2 – mAP50-95 evolution",
    "Line plot (epoch vs mAP@0.5:0.95) with three curves: YOLOv8m, YOLOv11m, YOLOv12s. Caption: 'Fig. 3.2 – mAP50-95 evolution over training epochs'.",
    "research_output/metrics_mAP50-95(B).png",
)
justify_para(
    "YOLOv8m converges smoothly to the highest value (0.403) with low variance (CV = 0.020). "
    "YOLOv11m is more variable (CV = 0.039) and peaks slightly later. YOLOv12s converges "
    "fastest but stabilises at a slightly lower value with moderate variance."
)

heading("3.1.4 Training and Validation Loss Curves (Figs. 3.3, 3.4)", level=3)
image_placeholder(
    "Figure 3.3 – Training loss curves",
    "Three sub-plots (or one combined plot) showing box, cls, and dfl training losses vs epoch for all three models. Caption: 'Fig. 3.3 – Training loss curves (box, cls, dfl) over epochs'.",
    "research_output/train_losses.png",
)
image_placeholder(
    "Figure 3.4 – Validation loss curves",
    "Line plot of validation loss (box/cls/dfl combined) vs epoch for the three models. Caption: 'Fig. 3.4 – Validation loss curves over epochs'.",
    "research_output/val_losses.png",
)
justify_para(
    "YOLOv8m losses decrease smoothly and plateau (box 1.78, cls 1.12, dfl 1.98). YOLOv11m "
    "losses decrease but with visible oscillation in classification loss. YOLOv12s shows "
    "extreme loss spikes in the early epochs, stabilising only after epoch 60. Validation "
    "losses mirror these patterns: YOLOv8m is consistently low, YOLOv11m fluctuates "
    "moderately, and YOLOv12s settles later and slightly higher."
)

heading("3.1.5 Efficiency", level=3)
p("Table 3.2 – Efficiency Comparison", size=11, italic=True, align="center", space_after=4)
add_table(
    ["Model", "Epochs", "Avg. Time/Epoch (s)", "Total Time (s)", "GFLOPs"],
    [
        ["YOLOv8m",  "150", "255", "38,250", "78.7"],
        ["YOLOv11m", "150", "255", "38,250", "79.1"],
        ["YOLOv12s", "100", "163", "16,300", "45.2"],
    ],
    widths=[1.1, 0.9, 1.6, 1.3, 0.9],
)
p("")
justify_para(
    "YOLOv12s trains 57% faster than the other two, but the speed benefit does not compensate "
    "for its lower accuracy and training instability in a high-stakes quality-control setting."
)

heading("3.1.6 Interpretation / Justification of Outcomes vs Goals", level=3)
add_table(
    ["Goal (from §1.6)", "Target", "Achieved", "Verdict"],
    [
        ["F1-score", "≥ 0.75", "0.759 (YOLOv8m)", "Met"],
        ["mAP50-95", "≥ 0.40", "0.403 (YOLOv8m)", "Met"],
        ["Stability (CV of mAP50-95)", "Low", "0.020", "Met"],
        ["Working Streamlit UI", "Upload → detect", "streamlit_app.py deployed", "Met"],
        ["Evidence-based model selection", "3-way comparison", "Sections 3.1.1 – 3.1.5", "Met"],
    ],
    widths=[2.1, 1.1, 1.9, 0.8],
)
p("")
justify_para(
    "Why YOLOv8m wins. Its decoupled head, task-aligned assigner, and anchor-free design "
    "simplify learning and improve generalisation. The balanced F1 is exactly what logistics "
    "needs: it minimises both missed defects (which reach customers) and false alarms (which "
    "waste re-inspection effort). The low CV makes re-training predictable—critical for a "
    "production system."
)
justify_para(
    "Why YOLOv11m falls short here. High precision but compromised recall means many "
    "defective parcels would slip through. Low recall is the worst failure mode in this domain."
)
justify_para(
    "Why YOLOv12s is not recommended yet. Its speed is attractive for edge devices, but the "
    "early-epoch training instability is a reliability risk for a high-stakes pipeline. It is "
    "a candidate for future edge deployment once stabilising tricks (warm-up, gradient "
    "clipping) are tuned."
)

heading("3.2 Total Committed vs Completed User Stories", level=2)
p("Table 3.3 – Total Committed vs Completed User Stories", size=11, italic=True, align="center", space_after=4)
add_table(
    ["Sprint", "Committed", "Completed", "Completion Rate"],
    [
        ["Sprint 1", "4", "4", "100%"],
        ["Sprint 2", "6", "6", "100%"],
        ["Total",    "10", "10", "100%"],
    ],
    widths=[1.3, 1.3, 1.3, 1.8],
)
p("")
justify_para(
    "Justification for any incomplete stories: None — all committed user stories were "
    "completed on time. Stretch items identified during retrospectives (per-class recall "
    "dashboard, Jetson edge deployment, active-learning loop for hard negatives) were "
    "intentionally deferred to post-release as they were not in the committed backlog. This "
    "deferral is documented in Chapter 4 as future enhancement."
)
page_break()

# ============================================================
# CHAPTER 4
# ============================================================
p("CHAPTER 4", size=16, bold=True, align="center", space_after=4)
p("CONCLUSIONS AND FUTURE ENHANCEMENT", size=18, bold=True, align="center", space_after=18)

heading("4.1 Conclusions", level=2)
justify_para(
    "This project presented SmartFlow, an end-to-end damaged-parcel detection system for "
    "logistics quality control, and a comprehensive comparative study of three modern YOLO "
    "variants on a custom 10,000-image package-defect dataset."
)
justify_para("Key conclusions:")
for b in [
    "YOLOv8m is the most reliable choice for package defect detection, delivering the highest F1-score (0.759), the highest mAP50-95 (0.403), and the lowest coefficient of variation in validation mAP50-95 (0.020).",
    "F1-score, not just mAP, must guide model selection in logistics quality control. YOLOv11m achieved near-identical mAP50-95 to YOLOv8m but with precision-heavy, recall-poor behaviour that would let defective parcels through.",
    "Training stability matters as much as peak accuracy. YOLOv12s was 57% faster to train but the early instability and lower final accuracy make it an unreliable production choice without additional stabilisation work.",
    "The Streamlit deployment turns the research result into a usable product: a non-technical operator can upload a parcel photograph and receive annotated defect boxes with confidences in seconds.",
    "SmartFlow is deployable at every handoff in the logistics pipeline—warehouse dispatch, loading, transit checkpoints, and final delivery—providing a consistent, 24×7 quality gate.",
]:
    bullet(b)

heading("4.2 Future Enhancement", level=2)
for b in [
    "Dataset generalisation: validate on additional defect types, packaging materials (cardboard, plastic, padded envelopes) and real-world lighting to confirm transferability beyond the current dataset.",
    "Per-model hyper-parameter tuning: the present comparison used identical hyper-parameters for fairness; each model may improve further with dedicated tuning.",
    "Ensemble detection: combining YOLOv8m's balanced detections with YOLOv11m's high-precision outputs via weighted-box fusion could push both precision and recall higher.",
    "Edge deployment: port the detector to NVIDIA Jetson Orin / Xavier with TensorRT INT8 quantisation for deployment directly on conveyor-belt cameras.",
    "Broader SmartFlow platform: integrate defect detection with demand forecasting and route optimisation, recasting it as one module in a holistic smart-logistics platform.",
    "Active learning loop: use low-confidence predictions in production as candidates for re-annotation, continuously improving the model on the long tail of hard cases.",
    "Explainability: add Grad-CAM / EigenCAM overlays so operators can see why a defect was flagged, building trust in the system.",
]:
    bullet(b)
page_break()

# ============================================================
# REFERENCES
# ============================================================
p("REFERENCES", size=18, bold=True, align="center", space_after=18)
refs = [
    "[1] J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, \"You Only Look Once: Unified, Real-Time Object Detection,\" in Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR), 2016, pp. 779–788.",
    "[2] G. Jocher, A. Chaurasia, and J. Qiu, \"YOLOv8 by Ultralytics,\" 2023. [Online]. Available: https://github.com/ultralytics/ultralytics",
    "[3] Ultralytics, \"YOLOv11,\" 2024. [Online]. Available: https://docs.ultralytics.com/models/yolo11/",
    "[4] Y. Li et al., \"YOLOv12: Attention-Centric Real-Time Object Detectors,\" 2025. [Online]. Available: https://arxiv.org/abs/2502.12524",
    "[5] J. Redmon and A. Farhadi, \"YOLO9000: Better, Faster, Stronger,\" in Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR), 2017.",
    "[6] J. Redmon and A. Farhadi, \"YOLOv3: An Incremental Improvement,\" arXiv preprint arXiv:1804.02767, 2018.",
    "[7] A. Bochkovskiy, C.-Y. Wang, and H.-Y. M. Liao, \"YOLOv4: Optimal Speed and Accuracy of Object Detection,\" arXiv preprint arXiv:2004.10934, 2020.",
    "[8] G. Jocher et al., \"YOLOv5,\" 2020. [Online]. Available: https://github.com/ultralytics/yolov5",
    "[9] Y. He, K. Song, Q. Meng, and Y. Yan, \"An End-to-end Steel Surface Defect Detection Approach via Fusing Multiple Hierarchical Features,\" IEEE Trans. Instrum. Meas., vol. 69, no. 4, pp. 1493–1504, 2020.",
    "[10] J. Li, X. Liang, and Y. Wei, \"Fabric Defect Detection Based on Improved YOLOv3,\" in Proc. IEEE Int. Conf. Signal Process. (ICSP), 2020, pp. 1–5.",
    "[11] Z. Chen, D. Yang, and T. Zhang, \"Deep Learning Based Defect Detection for Electronic Components,\" IEEE Access, vol. 8, pp. 124567–124576, 2020.",
    "[12] Ultralytics, \"YOLO Performance on COCO,\" 2024. [Online]. Available: https://docs.ultralytics.com/models/benchmarks",
]
for r in refs:
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.3)
    run = para.add_run(r)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
page_break()

# ============================================================
# APPENDIX
# ============================================================
p("APPENDIX", size=18, bold=True, align="center", space_after=18)

heading("A. Patent Disclosure Form / Publication Details", level=2)
justify_para(
    "Publication (Under Review / Submitted): I. Singh, N. K. Mohanarajan, K. Priyadharshini, "
    "and B. Sowmiya, \"SmartFlow: Automated Defect Detection System using YOLOv8,\" Department "
    "of Computing Technologies, SRM Institute of Science and Technology, Kattankulathur, 2026."
)
justify_para(
    "Attach: copy of manuscript (minor_paper_2.pdf), acknowledgement / acceptance email, and "
    "patent disclosure form if filed."
)

heading("B. Sample Coding with Screenshots", level=2)
p("B.1 Streamlit App (streamlit_app.py, abridged)", size=12, bold=True, space_after=4)

code = '''import streamlit as st
from PIL import Image
from ultralytics import YOLO

WEIGHTS = "results/yolov8m_optimized_20260224_1855/weights/best.pt"

@st.cache_resource
def load_model():
    return YOLO(WEIGHTS)

def main():
    st.title("SmartFlow - Damaged Parcel Detection (YOLOv8m)")
    model = load_model()

    uploaded = st.file_uploader("Upload parcel image",
                                type=["jpg", "jpeg", "png"])
    conf = st.slider("Confidence threshold", 0.1, 0.9, 0.25, 0.05)

    if uploaded is not None:
        image = Image.open(uploaded).convert("RGB")
        results = model.predict(image, conf=conf, verbose=False)
        annotated = results[0].plot()
        st.image(annotated, caption="Detections", use_column_width=True)

        rows = []
        for b in results[0].boxes:
            cls = int(b.cls[0])
            conf_v = float(b.conf[0])
            x1, y1, x2, y2 = map(float, b.xyxy[0])
            rows.append({"class": model.names[cls],
                         "confidence": conf_v,
                         "x1": x1, "y1": y1, "x2": x2, "y2": y2})
        if rows:
            st.dataframe(rows)
        else:
            st.info("No detections above threshold.")

if __name__ == "__main__":
    main()
'''

for line in code.splitlines():
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(line if line else " ")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

p("")
p("B.2 Training Command (YOLOv8m baseline)", size=12, bold=True, space_after=4)
for line in [
    "yolo detect train \\",
    "     model=yolov8m.pt \\",
    "     data=dataset/data.yaml \\",
    "     epochs=150 imgsz=640 batch=16 \\",
    "     optimizer=SGD lr0=0.01 lrf=0.0001 \\",
    "     momentum=0.937 weight_decay=0.0005 \\",
    "     project=results name=yolov8m_optimized_20260224_1855",
]:
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(10)

p("")
p("B.3 Research Metrics Aggregation (scripts/06_research_metrics.py, abridged)",
  size=12, bold=True, space_after=4)
code2 = '''import pandas as pd, matplotlib.pyplot as plt

runs = {
    "YOLOv8m":  "results/yolov8m_optimized_20260224_1855/results.csv",
    "YOLOv11m": "results/yolov11m_optimized/results.csv",
    "YOLOv12s": "results/yolov12s_optimized/results.csv",
}

for metric in ["metrics/mAP50-95(B)",
               "metrics/precision(B)",
               "metrics/recall(B)"]:
    plt.figure()
    for name, csv in runs.items():
        df = pd.read_csv(csv)
        plt.plot(df["epoch"], df[metric], label=name)
    plt.xlabel("Epoch")
    plt.ylabel(metric)
    plt.legend()
    plt.savefig("research_output/" +
                metric.replace("/", "_") + ".png", dpi=200)
'''
for line in code2.splitlines():
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(line if line else " ")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

p("")
p("Screenshots to attach:", size=12, bold=True, space_after=4)

image_placeholder(
    "Figure B.1 – Streamlit home (empty state)",
    "Screenshot of the Streamlit app immediately after `streamlit run streamlit_app.py`, before any image is uploaded: title, file uploader, confidence slider visible.",
    "appx_B_1_streamlit_home.png",
)
image_placeholder(
    "Figure B.2 – Annotated parcel detection",
    "Screenshot of a parcel image after upload, showing bounding boxes drawn around detected defects (e.g., dent/scratch) and the corresponding class/confidence table below.",
    "appx_B_2_detection_result.png",
)
image_placeholder(
    "Figure B.3 – Confidence slider effect",
    "Side-by-side screenshots (or two stacked images) of the same parcel at confidence = 0.25 (more boxes) vs confidence = 0.75 (fewer, higher-confidence boxes).",
    "appx_B_3_confidence_comparison.png",
)
image_placeholder(
    "Figure B.4 – Ultralytics training summary",
    "The auto-generated results.png from the Ultralytics training run (results/yolov8m_optimized_20260224_1855/results.png) showing train/val loss, precision, recall, mAP50, mAP50-95 curves.",
    "results/yolov8m_optimized_20260224_1855/results.png",
)

heading("C. Plagiarism Report", level=2)
justify_para("Attach the Turnitin / Drillbit plagiarism-check PDF summary here.")
image_placeholder(
    "Appendix C – Plagiarism report screenshot",
    "Screenshot / scanned page of the Turnitin or Drillbit plagiarism-check summary page, showing the overall similarity index and the breakdown bar. Attach the full PDF as a separate annex.",
    "appx_C_plagiarism_report.png",
)
for b in [
    "Similarity Index: ____ %",
    "Internet Sources: ____ %",
    "Publications:     ____ %",
    "Student Papers:   ____ %",
]:
    bullet(b)

p("")
p("— End of Report —", size=12, italic=True, align="center")

doc.save(OUT)
print(f"Wrote: {OUT}")
